"""
Service layer for parsing documents in stream.
1. Parse the document into chunks with semantic segmentation by Tika.
2. Splict chunks into paragraphs.
3. Split paragraphs into sentences.
4. Split sentences into words. Chinese depends on Jieba, HanLP; English depends on NLTK.
5. Return the parsed result in DocumentVector format.
"""

import logging
from pydoc import Doc
import re
import os
import psutil
import gc
from typing import List, Optional, BinaryIO
from io import BytesIO
from dotenv import load_dotenv
import jieba
import nltk
from nltk.tokenize import sent_tokenize, word_tokenize
from pypdf import PdfReader
import docx 
from sqlalchemy.orm import Session
from app.models import DocumentVector
from app.repositories.document_vector_repository import create_document_vector, batch_create_document_vectors
from app.utils.logging import get_logger

# Initialize logger
logger = get_logger(__name__)

# Load environment variables
load_dotenv()

# Download required NLTK data (run once)
try:
    nltk.data.find('tokenizers/punkt')
except LookupError:
    nltk.download('punkt', quiet=True)

# Load parsing configuration from environment variables
CHUNK_SIZE = int(os.getenv("FILE_PARSING_CHUNK_SIZE", "1000"))
PARENT_CHUNK_SIZE = int(os.getenv("FILE_PARSING_PARENT_CHUNK_SIZE", "1048576"))
BUFFER_SIZE = int(os.getenv("FILE_PARSING_BUFFER_SIZE", "8192"))
MAX_MEMORY_THRESHOLD = float(os.getenv("FILE_PARSING_MAX_MEMORY_THRESHOLD", "0.8"))

logger.info(
    f"ParseService configuration loaded: "
    f"CHUNK_SIZE={CHUNK_SIZE}, "
    f"PARENT_CHUNK_SIZE={PARENT_CHUNK_SIZE}, "
    f"BUFFER_SIZE={BUFFER_SIZE}, "
    f"MAX_MEMORY_THRESHOLD={MAX_MEMORY_THRESHOLD}"
)


# ============================================================
# Streaming PDF Parser (by page)
# ============================================================
class PdfTextIterator:
    def __init__(self, file_stream:BinaryIO):
        self.reader = PdfReader(file_stream)
        self.pages = self.reader.pages
        self.index = 0
    
    def __iter__(self):
        return self
    
    def __next__(self):
        if self.index >= len(self.pages):
            raise StopIteration
        
        text = self.pages[self.index].extract_text() or ""
        self.index += 1
        return text

# ============================================================
# Streaming DOCX Parser (by paragraph)
# ============================================================
class DocxTextIterator:
    def __init__(self, file_stream:BinaryIO):
        self.doc = docx.Document(file_stream)
        self.paragraphs = self.doc.paragraphs
        self.index = 0
    
    def __iter__(self):
        return self

    def __next__(self):
        if self.index >= len(self.paragraphs):
            raise StopIteration
        
        text = self.paragraphs[self.index].text or ""
        self.index += 1
        return text

# ============================================================
# Streaming TXT Parser (by line)
# ============================================================
class PlainTextIterator:
    def __init__(self, file_stream:BinaryIO, buffer_size: int = 8192):
        self.stream = file_stream
        self.buffer_size = buffer_size
    
    def __iter__(self):
        return self

    def __next__(self):
        data = self.stream.read(self.buffer_size)
        if not data:
            raise StopIteration
        return data.decode("utf-8", errors="ignore") 

class ParseService:
    """
    Document parsing service with streaming support to avoid OOM.
    Uses parent-child chunking strategy.
    """
    def __init__(self, chunk_size: Optional[int] = None, parent_chunk_size: Optional[int] = None, buffer_size: Optional[int] = None, max_memory_threshold: Optional[float] = None):
        self.chunk_size = chunk_size or CHUNK_SIZE
        self.parent_chunk_size = parent_chunk_size or PARENT_CHUNK_SIZE
        self.buffer_size = buffer_size or BUFFER_SIZE
        self.max_memory_threshold = max_memory_threshold or MAX_MEMORY_THRESHOLD

        # Initialize jieba for Chinese tokenization
        jieba.initialize()
        
        logger.info(
            f"ParseService instance initialized with chunk_size={self.chunk_size}, "
            f"parent_chunk_size={self.parent_chunk_size}, "
            f"buffer_size={self.buffer_size}, "
            f"max_memory_threshold={self.max_memory_threshold}"
        )

    def check_memory(self):
        """
        Check if the memory usage is too high and trigger garbage collection.
        """
        vm = psutil.virtual_memory()
        proc = psutil.Process()
        mem_usage_ratio = proc.memory_info().rss / vm.total

        if mem_usage_ratio > self.max_memory_threshold:
            gc.collect()
            if mem_usage_ratio > self.max_memory_threshold:
                raise MemoryError("Memory usage exceeds the threshold.")


    def _detect_file_type(self, file_name: str) -> str:
        """
        Detect the file type from the filename.
        """
        file_name = file_name.lower()
        if file_name.endswith(".pdf"):
            return "pdf"
        if file_name.endswith(".docx"):
            return "docx"
        if file_name.endswith(".txt") or file_name.endswith(".md"):
            return "text"
        return "unknown"

    def get_iterator(self, file_name: str, file_stream: BinaryIO):
        """
        Choose appropriate iterator based on the file type.
        """
        ftype = self._detect_file_type(file_name)

        if ftype == "pdf":
            return PdfTextIterator(file_stream)
        if ftype == "docx":
            return DocxTextIterator(file_stream)
        if ftype == "text":
            return PlainTextIterator(file_stream, self.buffer_size)

        # fallback
        return PlainTextIterator(file_stream, self.buffer_size)


    def parse_and_save(self, file_md5: str, file_name: str, file_stream: BinaryIO, db: Session):
        """
        Process the file stream in parent-child chunks and save the parsed result to the database.
        """
        parent_buffer = []
        parent_length = 0
        chunk_id = 0

        iterator = self.get_iterator(file_name, file_stream)

        for block in iterator:
            self.check_memory()
            
            parent_buffer.append(block)
            parent_length += len(block)

            if parent_length >= self.parent_chunk_size:
                chunk_id = self.process_parent_chunk(file_md5, parent_buffer, chunk_id, db)
                parent_buffer = []
                parent_length = 0
            
        if parent_buffer:
            chunk_id = self.process_parent_chunk(file_md5, parent_buffer, chunk_id, db)
        
        return chunk_id
    
    def process_parent_chunk(self, file_md5: str, parent_buffer: List[str], start_chunk_id: int, db: Session):
        """
        Process the parent chunk and save the parsed result to the database.
        """
        parent_text = "".join(parent_buffer)
        chunks = self.split_text_into_chunks_with_semantics(parent_text)

        vectors = []
        chunk_id = start_chunk_id

        for chunk in chunks:
            chunk_id += 1
            vectors.append(DocumentVector(file_md5=file_md5, chunk_id=chunk_id, text_content=chunk, model_version="default"))
        
        batch_create_document_vectors(db, vectors)
        return chunk_id
    
    def split_text_into_chunks_with_semantics(self, text: str) -> List[str]:
        """
        Split the text into chunks with semantic segmentation.
        """
        chunks = []
        paragraphs = re.split(r"\n\n+", text)

        current_chunk = []

        for paragraph in paragraphs:
            paragraph = paragraph.strip()
            if not paragraph:
                continue
            
            if len(paragraph) >= self.chunk_size:
                if current_chunk:
                    chunks.append("".join(current_chunk).strip())
                    current_chunk = []
                chunks.extend(self.split_paragraph_into_sentences(paragraph))
                continue

            if sum(len(x) for x in current_chunk) + len(paragraph) > self.chunk_size:
                if current_chunk:
                    chunks.append("".join(current_chunk).strip())
                current_chunk = [paragraph]
            else:
                if current_chunk:
                    current_chunk.append("\n\n")
                current_chunk.append(paragraph)
        
        if current_chunk:
            chunks.append("".join(current_chunk).strip())

        return chunks
    
    def split_paragraph_into_sentences(self, paragraph: str) -> List[str]:
        """
        Split the paragraph into sentences.
        """
        chunks = []
        sentences = re.split(r"(?<=[。！？；])|(?<=[.!?;])\s+", paragraph)

        current = []

        for sentence in sentences:
            sentence = sentence.strip()
            if not sentence:
                continue

            # 句子太长 → 按词语拆
            if len(sentence) > self.chunk_size:
                if current:
                    chunks.append("".join(current).strip())
                    current = []
                chunks.extend(self.split_long_sentence(sentence))
                continue

            # 能否拼进当前 chunk？
            if sum(len(x) for x in current) + len(sentence) > self.chunk_size:
                chunks.append("".join(current).strip())
                current = [sentence]
            else:
                current.append(sentence)

        if current:
            chunks.append("".join(current).strip())

        return chunks

    def split_long_sentence(self, sentence: str) -> List[str]:
        """
        Split the long sentence into chunks.
        """
        chunks = []

        ascii_ratio = sum(c.isascii() for c in sentence) / max(len(sentence), 1)
        if ascii_ratio > 0.7:
            words = word_tokenize(sentence)  # 英文
        else:
            words = list(jieba.cut(sentence))  # 中文

        current = []
        current_len = 0

        for w in words:
            if current_len + len(w) > self.chunk_size and current:
                chunks.append("".join(current))
                current = []
                current_len = 0

            current.append(w)
            current_len += len(w)

        if current:
            chunks.append("".join(current))

        return chunks



        
    
